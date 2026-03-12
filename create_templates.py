"""
Run this script once to generate the placeholder .docx templates.
Replace the placeholder text with your actual NDA and Contract content.

Usage:
    python create_templates.py
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path

TEMPLATES_DIR = Path(__file__).parent / "templates"
TEMPLATES_DIR.mkdir(exist_ok=True)


def make_nda_template():
    doc = Document()

    # Title
    title = doc.add_heading("NON-DISCLOSURE AGREEMENT", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("")

    doc.add_paragraph(
        "This Non-Disclosure Agreement (\"Agreement\") is entered into as of {{date}} "
        "between the Company and the following individual:"
    )

    doc.add_paragraph("")

    # Employee info table
    table = doc.add_table(rows=3, cols=2)
    table.style = "Table Grid"
    table.cell(0, 0).text = "Full Name"
    table.cell(0, 1).text = "{{full_name}}"
    table.cell(1, 0).text = "ID / Passport Number"
    table.cell(1, 1).text = "{{id_number}}"
    table.cell(2, 0).text = "Position"
    table.cell(2, 1).text = "{{position}}"

    doc.add_paragraph("")

    doc.add_heading("1. Confidential Information", level=2)
    doc.add_paragraph(
        "\"Confidential Information\" means any information disclosed by the Company to "
        "the Employee, either directly or indirectly, in writing, orally, or by inspection "
        "of tangible objects, including without limitation technical data, trade secrets, "
        "know-how, research, product plans, products, services, customers, markets, "
        "software, developments, inventions, processes, formulas, technology, designs, "
        "drawings, engineering, hardware configuration information, marketing, finances, "
        "or other business information."
    )

    doc.add_heading("2. Obligations", level=2)
    doc.add_paragraph(
        "The Employee agrees to: (a) hold all Confidential Information in strict confidence; "
        "(b) not disclose Confidential Information to any third parties without prior written "
        "consent of the Company; (c) use Confidential Information solely for the purpose of "
        "performing duties for the Company."
    )

    doc.add_heading("3. Term", level=2)
    doc.add_paragraph(
        "This Agreement shall remain in effect for a period of three (3) years from the "
        "date of signing, and shall survive termination of employment."
    )

    doc.add_heading("4. Governing Law", level=2)
    doc.add_paragraph(
        "This Agreement shall be governed by applicable law."
    )

    doc.add_paragraph("")
    doc.add_paragraph("")

    # Signature block
    sig_table = doc.add_table(rows=3, cols=2)
    sig_table.cell(0, 0).text = "COMPANY"
    sig_table.cell(0, 1).text = "EMPLOYEE"
    sig_table.cell(1, 0).text = "Signature: ___________________"
    sig_table.cell(1, 1).text = "Signature: ___________________"
    sig_table.cell(2, 0).text = "Date: ___________________"
    sig_table.cell(2, 1).text = "Date: ___________________"

    path = TEMPLATES_DIR / "nda_template.docx"
    doc.save(str(path))
    print(f"Created: {path}")


def make_contract_template():
    doc = Document()

    title = doc.add_heading("SERVICE CONTRACT", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("")
    doc.add_paragraph(
        "This Service Contract (\"Contract\") is entered into as of {{date}} between "
        "the Company and the Contractor specified below."
    )

    doc.add_paragraph("")

    table = doc.add_table(rows=6, cols=2)
    table.style = "Table Grid"
    rows = [
        ("Full Name", "{{full_name}}"),
        ("Position", "{{position}}"),
        ("IE Registration Number", "{{ie_number}}"),
        ("Bank Name", "{{bank_name}}"),
        ("IBAN", "{{iban}}"),
        ("Monthly Fee", "{{salary}}"),
    ]
    for i, (label, value) in enumerate(rows):
        table.cell(i, 0).text = label
        table.cell(i, 1).text = value

    doc.add_paragraph("")

    doc.add_heading("1. Scope of Services", level=2)
    doc.add_paragraph(
        "The Contractor agrees to perform services as described in the attached Scope of Work "
        "document and any additional tasks mutually agreed upon in writing."
    )

    doc.add_heading("2. Compensation", level=2)
    doc.add_paragraph(
        "The Company shall pay the Contractor a monthly fee of {{salary}} upon submission "
        "of a monthly work report and invoice."
    )

    doc.add_heading("3. Term", level=2)
    doc.add_paragraph(
        "This Contract is effective from {{date}} and shall continue until terminated "
        "by either party with 30 days written notice."
    )

    doc.add_heading("4. Intellectual Property", level=2)
    doc.add_paragraph(
        "All work product, inventions, and deliverables created by the Contractor in "
        "connection with the services shall be the sole property of the Company."
    )

    doc.add_heading("5. Confidentiality", level=2)
    doc.add_paragraph(
        "The Contractor agrees to keep all Company information confidential, consistent "
        "with any separate NDA in force between the parties."
    )

    doc.add_paragraph("")
    doc.add_paragraph("")

    sig_table = doc.add_table(rows=3, cols=2)
    sig_table.cell(0, 0).text = "COMPANY"
    sig_table.cell(0, 1).text = "CONTRACTOR: {{full_name}}"
    sig_table.cell(1, 0).text = "Signature: ___________________"
    sig_table.cell(1, 1).text = "Signature: ___________________"
    sig_table.cell(2, 0).text = "Date: ___________________"
    sig_table.cell(2, 1).text = "Date: ___________________"

    path = TEMPLATES_DIR / "contract_template.docx"
    doc.save(str(path))
    print(f"Created: {path}")


if __name__ == "__main__":
    make_nda_template()
    make_contract_template()
    print("Done. Edit the .docx files in templates/ to match your actual legal documents.")
