import io
from datetime import date
from pathlib import Path

from docx import Document
from fastapi import FastAPI, HTTPException
from fastapi.responses import Response
from pydantic import BaseModel

app = FastAPI(title="NVS Doc Generator", version="1.0.0")

TEMPLATES_DIR = Path(__file__).parent / "templates"


# ─── Request Models ──────────────────────────────────────────


class NDARequest(BaseModel):
    full_name: str
    id_number: str
    position: str
    date: str = str(date.today())


class ContractRequest(BaseModel):
    full_name: str
    position: str
    ie_number: str
    bank_name: str
    iban: str
    salary: str
    date: str = str(date.today())


# ─── Helpers ─────────────────────────────────────────────────


def fill_template(template_path: Path, variables: dict) -> bytes:
    """Replace {{variable}} placeholders in all paragraph runs of a docx, return bytes."""
    doc = Document(str(template_path))

    def replace_in_text(text: str) -> str:
        for key, value in variables.items():
            text = text.replace("{{" + key + "}}", str(value))
        return text

    for para in doc.paragraphs:
        for run in para.runs:
            run.text = replace_in_text(run.text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.text = replace_in_text(run.text)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ─── Endpoints ───────────────────────────────────────────────


@app.get("/health")
def health():
    return {"status": "ok"}


@app.post("/generate-nda", summary="Generate NDA .docx from template")
def generate_nda(req: NDARequest):
    template_path = TEMPLATES_DIR / "nda_template.docx"
    if not template_path.exists():
        raise HTTPException(status_code=500, detail="NDA template not found")

    variables = {
        "full_name": req.full_name,
        "id_number": req.id_number,
        "position": req.position,
        "date": req.date,
    }

    docx_bytes = fill_template(template_path, variables)
    filename = f"NDA_{req.full_name.replace(' ', '_')}.docx"
    return Response(
        content=docx_bytes,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@app.post("/generate-contract", summary="Generate Contract .docx from template")
def generate_contract(req: ContractRequest):
    template_path = TEMPLATES_DIR / "contract_template.docx"
    if not template_path.exists():
        raise HTTPException(status_code=500, detail="Contract template not found")

    variables = {
        "full_name": req.full_name,
        "position": req.position,
        "ie_number": req.ie_number,
        "bank_name": req.bank_name,
        "iban": req.iban,
        "salary": req.salary,
        "date": req.date,
    }

    docx_bytes = fill_template(template_path, variables)
    filename = f"Contract_{req.full_name.replace(' ', '_')}.docx"
    return Response(
        content=docx_bytes,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )
